"""
Module xuất báo cáo: .docx và .pdf
"""

import os
import io
from datetime import datetime, timezone

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table as RLTable,
    TableStyle, PageBreak, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# ═══════════════════════════════════════════════════════════════
# XUẤT DOCX
# ═══════════════════════════════════════════════════════════════

def export_docx(analysis: dict, filename: str = "report.docx") -> str:
    """Xuất báo cáo phân tích ra file .docx"""

    doc = DocxDocument()

    # Thiết lập style
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)  # Font chữ đỏ

    # --- TRANG BÌA ---
    title = doc.add_heading("", level=0)
    run = title.add_run("BÁO CÁO THU THẬP VÀ PHÂN TÍCH\nTHÔNG TIN QUỐC TẾ")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Phục vụ an ninh đối ngoại Việt Nam")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now = datetime.now(timezone.utc)
    run = meta.add_run(f"\nNgày tạo: {now.strftime('%d/%m/%Y %H:%M UTC')}\n"
                       f"Tác giả: Vietmetric\n"
                       f"Tổng bài viết: {analysis.get('total', 0)}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()

    # --- MỤC LỤC ---
    doc.add_heading("Mục lục", level=1).runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    toc_items = [
        "1. Tóm tắt",
        "2. Đánh giá mức độ đe dọa",
        "3. Chủ đề nổi bật",
        "4. Thực thể địa chính trị",
        "5. Tin nghiêm trọng và quan trọng",
        "6. Danh sách toàn bộ nguồn tin",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()

    # --- 1. TÓM TẮT ---
    doc.add_heading("1. Tóm tắt", level=1).runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    summary = analysis.get("summary", "Không có dữ liệu.")
    for line in summary.split("\n"):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # --- 2. ĐÁNH GIÁ MỨC ĐỘ ---
    doc.add_heading("2. Đánh giá mức độ đe dọa", level=1).runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    threat = analysis.get("threat_assessment", {})

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Mức độ"
    hdr[1].text = "Số lượng"
    hdr[2].text = "Tỷ lệ"

    total = max(sum(threat.values()), 1)
    levels_vi = {
        "critical": "Nghiêm trọng",
        "high": "Cao",
        "medium": "Trung bình",
        "low": "Thấp"
    }
    for level_key in ["critical", "high", "medium", "low"]:
        count = threat.get(level_key, 0)
        row = table.add_row().cells
        row[0].text = levels_vi.get(level_key, level_key)
        row[1].text = str(count)
        row[2].text = f"{count/total*100:.1f}%"

    # --- 3. CHỦ ĐỀ NỔI BẬT ---
    doc.add_heading("3. Chủ đề nổi bật", level=1).runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    topics = analysis.get("key_topics", [])
    if topics:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Từ khóa"
        hdr[1].text = "Tần suất"
        for topic in topics[:15]:
            row = table.add_row().cells
            row[0].text = topic["keyword"]
            row[1].text = str(topic["count"])
    else:
        p = doc.add_paragraph("Không phát hiện chủ đề nổi bật.")
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # --- 4. THỰC THỂ ĐỊA CHÍNH TRỊ ---
    doc.add_heading("4. Thực thể địa chính trị", level=1).runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    entities = analysis.get("geographic_focus", [])
    if entities:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Quốc gia / Tổ chức"
        hdr[1].text = "Số lần xuất hiện"
        for entity in entities[:15]:
            row = table.add_row().cells
            row[0].text = entity["entity"]
            row[1].text = str(entity["count"])

    # --- 5. TIN NGHIÊM TRỌNG ---
    doc.add_heading("5. Tin nghiêm trọng và quan trọng", level=1).runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    articles = analysis.get("analyzed_articles", [])
    important = [a for a in articles if a.get("relevance_level") in ("critical", "high")]

    for i, art in enumerate(important[:30]):
        level_vi = levels_vi.get(art.get("relevance_level", "low"), "Thấp")
        score = art.get("relevance_score", 0)

        doc.add_heading(f"{i+1}. {art.get('title', 'Không có tiêu đề')}", level=2).runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        info = doc.add_paragraph()
        run = info.add_run(f"Nguồn: {art.get('source', 'N/A')} | "
                          f"Ngôn ngữ: {art.get('source_lang_label', 'N/A')} | "
                          f"Mức độ: {level_vi} ({score} điểm)")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if art.get("summary"):
            p = doc.add_paragraph(art["summary"])
            p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        url_p = doc.add_paragraph()
        run = url_p.add_run(f"URL: {art.get('url', '')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)

        keywords = art.get("matched_keywords", [])
        if keywords:
            kw_p = doc.add_paragraph()
            run = kw_p.add_run(f"Từ khóa: {', '.join(keywords[:5])}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # --- 6. DANH SÁCH TOÀN BỘ ---
    doc.add_page_break()
    doc.add_heading("6. Danh sách toàn bộ nguồn tin", level=1).runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "STT"
    hdr[1].text = "Tiêu đề"
    hdr[2].text = "Nguồn"
    hdr[3].text = "Mức độ"
    hdr[4].text = "Điểm"

    for i, art in enumerate(articles[:100]):
        row = table.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = art.get("title", "")[:80]
        row[2].text = art.get("source", "")
        row[3].text = levels_vi.get(art.get("relevance_level", "low"), "Thấp")
        row[4].text = str(art.get("relevance_score", 0))

    # Footer
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Báo cáo tạo tự động bởi Vietmetric News Collector - Vietmetric\n"
                         f"Thời điểm: {now.strftime('%d/%m/%Y %H:%M UTC')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(filename)
    return filename


# ═══════════════════════════════════════════════════════════════
# XUẤT PDF
# ═══════════════════════════════════════════════════════════════

def export_pdf(analysis: dict, filename: str = "report.pdf") -> str:
    """Xuất báo cáo phân tích ra file .pdf"""

    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        name="VietTitle",
        parent=styles["Title"],
        fontSize=20,
        textColor=HexColor("#CC0000"),
        spaceAfter=20,
        alignment=TA_CENTER
    ))
    styles.add(ParagraphStyle(
        name="VietHeading",
        parent=styles["Heading1"],
        fontSize=14,
        textColor=HexColor("#CC0000"),
        spaceBefore=16,
        spaceAfter=8
    ))
    styles.add(ParagraphStyle(
        name="VietBody",
        parent=styles["Normal"],
        fontSize=10,
        textColor=HexColor("#CC0000"),
        spaceAfter=6,
        alignment=TA_JUSTIFY
    ))
    styles.add(ParagraphStyle(
        name="VietSmall",
        parent=styles["Normal"],
        fontSize=8,
        textColor=HexColor("#666666"),
        spaceAfter=4
    ))
    styles.add(ParagraphStyle(
        name="VietURL",
        parent=styles["Normal"],
        fontSize=8,
        textColor=HexColor("#0000CC"),
        spaceAfter=4
    ))

    story = []

    # Trang bìa
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph("BAO CAO THU THAP VA PHAN TICH<br/>THONG TIN QUOC TE", styles["VietTitle"]))
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("Phuc vu an ninh doi ngoai Viet Nam", styles["VietBody"]))
    story.append(Spacer(1, 2*cm))

    now = datetime.now(timezone.utc)
    meta_text = (f"Ngay tao: {now.strftime('%d/%m/%Y %H:%M UTC')}<br/>"
                 f"Tac gia: Vietmetric<br/>"
                 f"Tong bai viet: {analysis.get('total', 0)}")
    story.append(Paragraph(meta_text, styles["VietBody"]))
    story.append(PageBreak())

    # 1. Tóm tắt
    story.append(Paragraph("1. Tom tat", styles["VietHeading"]))
    summary = analysis.get("summary", "Khong co du lieu.")
    for line in summary.split("\n"):
        if line.strip():
            safe_line = line.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_line, styles["VietBody"]))

    # 2. Đánh giá mức độ
    story.append(Paragraph("2. Danh gia muc do de doa", styles["VietHeading"]))
    threat = analysis.get("threat_assessment", {})

    levels_vi = {
        "critical": "Nghiem trong",
        "high": "Cao",
        "medium": "Trung binh",
        "low": "Thap"
    }
    total = max(sum(threat.values()), 1)

    table_data = [["Muc do", "So luong", "Ty le"]]
    for level_key in ["critical", "high", "medium", "low"]:
        count = threat.get(level_key, 0)
        table_data.append([
            levels_vi.get(level_key, level_key),
            str(count),
            f"{count/total*100:.1f}%"
        ])

    t = RLTable(table_data, colWidths=[6*cm, 4*cm, 4*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#CC0000")),
        ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#FFF5F5"), HexColor("#FFFFFF")])
    ]))
    story.append(t)
    story.append(Spacer(1, 0.5*cm))

    # 3. Chủ đề nổi bật
    story.append(Paragraph("3. Chu de noi bat", styles["VietHeading"]))
    topics = analysis.get("key_topics", [])
    if topics:
        topic_data = [["Tu khoa", "Tan suat"]]
        for topic in topics[:15]:
            topic_data.append([topic["keyword"], str(topic["count"])])

        t = RLTable(topic_data, colWidths=[10*cm, 4*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#CC0000")),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
            ("ALIGN", (1, 0), (1, -1), "CENTER"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
        ]))
        story.append(t)

    # 5. Tin quan trọng
    story.append(PageBreak())
    story.append(Paragraph("4. Tin nghiem trong va quan trong", styles["VietHeading"]))
    articles = analysis.get("analyzed_articles", [])
    important = [a for a in articles if a.get("relevance_level") in ("critical", "high")]

    for i, art in enumerate(important[:20]):
        level_vi = levels_vi.get(art.get("relevance_level", "low"), "Thap")
        title_text = f"<b>{i+1}. {art.get('title', 'N/A')}</b>"
        title_text = title_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        title_text = f"<b>{i+1}.</b> " + art.get("title", "N/A").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        story.append(Paragraph(title_text, styles["VietBody"]))
        info = (f"Nguon: {art.get('source', 'N/A')} | "
                f"Ngon ngu: {art.get('source_lang_label', 'N/A')} | "
                f"Muc do: {level_vi} ({art.get('relevance_score', 0)} diem)")
        story.append(Paragraph(info, styles["VietSmall"]))

        if art.get("summary"):
            safe_summary = art["summary"][:300].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_summary, styles["VietSmall"]))

        url = art.get("url", "")
        if url:
            safe_url = url.replace("&", "&amp;")
            story.append(Paragraph(f"URL: {safe_url}", styles["VietURL"]))

        story.append(Spacer(1, 0.3*cm))

    # Footer
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph(
        f"Bao cao tao tu dong boi Vietmetric News Collector - Vietmetric | {now.strftime('%d/%m/%Y %H:%M UTC')}",
        styles["VietSmall"]
    ))

    doc.build(story)
    return filename
